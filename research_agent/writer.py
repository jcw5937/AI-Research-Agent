"""Word document generation for research notes."""

import os

from docx import Document
from docx.shared import Pt

from project import DEFAULT_NOTE_SECTIONS


def build_layer_document(
    layer_name: str,
    papers_notes: list[dict],
    output_dir: str,
    note_sections: list[str] | None = None,
) -> str:
    """Build a single Word document for a layer's papers and return the output path."""
    note_sections = note_sections or DEFAULT_NOTE_SECTIONS
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()
    doc.add_heading(layer_name, level=0)

    for paper in papers_notes:
        title = paper.get("title", "Untitled Paper")
        authors = paper.get("authors") or "Unknown authors"
        venue = paper.get("venue") or "Unknown venue"
        url = paper.get("url", "")
        verified = "verified" if paper.get("verified") else "unverified"
        peer_reviewed = "peer-reviewed" if paper.get("peer_reviewed") else "not peer-reviewed"
        sections = paper.get("sections", {})

        doc.add_heading(title, level=1)

        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"{authors} — {venue} | {verified} | {peer_reviewed}")
        meta_run.italic = True
        meta_run.font.size = Pt(10)

        if url:
            url_para = doc.add_paragraph()
            url_run = url_para.add_run(url)
            url_run.italic = True
            url_run.font.size = Pt(9)

        for section_name in note_sections:
            doc.add_heading(section_name, level=2)
            body_text = sections.get(section_name, "").strip() or "No information extracted."
            doc.add_paragraph(body_text)

    safe_name = "".join(c if c.isalnum() or c in (" ", "-", "_") else "_" for c in layer_name).strip()
    safe_name = safe_name.replace(" ", "_") or "layer"
    file_path = os.path.join(output_dir, f"{safe_name}.docx")
    doc.save(file_path)
    return file_path
